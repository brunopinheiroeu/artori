import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import json
from langchain.schema import Document
from langchain.document_loaders import (
    PyPDFLoader,
    TextLoader,
    DirectoryLoader
)
from langchain.document_loaders.base import BaseLoader
import docx
from rag_service import rag_service

# Configure logging
logger = logging.getLogger(__name__)

class DataIngestionPipeline:
    """Pipeline for ingesting various document types into the RAG system"""
    
    def __init__(self):
        self.supported_extensions = {'.pdf', '.txt', '.docx', '.json'}
        
    def load_document(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> List[Document]:
        """
        Load a single document from file path
        
        Args:
            file_path: Path to the document
            metadata: Additional metadata to attach to the document
            
        Returns:
            List of Document objects
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return []
        
        if file_path.suffix.lower() not in self.supported_extensions:
            logger.warning(f"Unsupported file type: {file_path.suffix}")
            return []
        
        try:
            documents = []
            base_metadata = metadata or {}
            base_metadata.update({
                'source': str(file_path),
                'filename': file_path.name,
                'file_type': file_path.suffix.lower()
            })
            
            # Ensure content_language is set (defaults to 'en' if not specified)
            if 'content_language' not in base_metadata:
                base_metadata['content_language'] = base_metadata.get('language', 'en')
            
            if file_path.suffix.lower() == '.pdf':
                documents = self._load_pdf(file_path, base_metadata)
            elif file_path.suffix.lower() == '.txt':
                documents = self._load_text(file_path, base_metadata)
            elif file_path.suffix.lower() == '.docx':
                documents = self._load_docx(file_path, base_metadata)
            elif file_path.suffix.lower() == '.json':
                documents = self._load_json(file_path, base_metadata)
            
            logger.info(f"✅ Loaded {len(documents)} documents from {file_path}")
            return documents
            
        except Exception as e:
            logger.error(f"❌ Failed to load document {file_path}: {e}")
            return []
    
    def _load_pdf(self, file_path: Path, metadata: Dict[str, Any]) -> List[Document]:
        """Load PDF document"""
        loader = PyPDFLoader(str(file_path))
        documents = loader.load()
        
        # Add metadata to each document
        for doc in documents:
            doc.metadata.update(metadata)
        
        return documents
    
    def _load_text(self, file_path: Path, metadata: Dict[str, Any]) -> List[Document]:
        """Load text document"""
        loader = TextLoader(str(file_path), encoding='utf-8')
        documents = loader.load()
        
        # Add metadata to each document
        for doc in documents:
            doc.metadata.update(metadata)
        
        return documents
    
    def _load_docx(self, file_path: Path, metadata: Dict[str, Any]) -> List[Document]:
        """Load DOCX document"""
        try:
            doc = docx.Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())
            
            full_text = '\n'.join(content)
            
            document = Document(
                page_content=full_text,
                metadata=metadata
            )
            
            return [document]
            
        except Exception as e:
            logger.error(f"Failed to load DOCX file {file_path}: {e}")
            return []
    
    def _load_json(self, file_path: Path, metadata: Dict[str, Any]) -> List[Document]:
        """Load JSON document (for structured educational content)"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            documents = []
            
            # Handle different JSON structures
            if isinstance(data, list):
                # List of items (e.g., questions, topics)
                for i, item in enumerate(data):
                    content = self._extract_content_from_json_item(item)
                    if content:
                        doc_metadata = metadata.copy()
                        doc_metadata.update({
                            'item_index': i,
                            'item_type': type(item).__name__
                        })
                        
                        # Add specific metadata from the item
                        if isinstance(item, dict):
                            for key in ['subject', 'topic', 'difficulty', 'language', 'content_language']:
                                if key in item:
                                    doc_metadata[key] = item[key]
                            
                            # Map 'language' to 'content_language' for backward compatibility
                            if 'language' in item and 'content_language' not in doc_metadata:
                                doc_metadata['content_language'] = item['language']
                        
                        documents.append(Document(
                            page_content=content,
                            metadata=doc_metadata
                        ))
            
            elif isinstance(data, dict):
                # Single structured document
                content = self._extract_content_from_json_item(data)
                if content:
                    doc_metadata = metadata.copy()
                    
                    # Add specific metadata from the data
                    for key in ['subject', 'topic', 'difficulty', 'language', 'content_language']:
                        if key in data:
                            doc_metadata[key] = data[key]
                    
                    # Map 'language' to 'content_language' for backward compatibility
                    if 'language' in data and 'content_language' not in doc_metadata:
                        doc_metadata['content_language'] = data['language']
                    
                    documents.append(Document(
                        page_content=content,
                        metadata=doc_metadata
                    ))
            
            return documents
            
        except Exception as e:
            logger.error(f"Failed to load JSON file {file_path}: {e}")
            return []
    
    def _extract_content_from_json_item(self, item: Any) -> str:
        """Extract meaningful content from a JSON item"""
        if isinstance(item, str):
            return item
        
        if isinstance(item, dict):
            content_parts = []
            
            # Common educational content fields
            content_fields = [
                'question', 'text', 'content', 'description', 
                'explanation', 'answer', 'solution', 'title'
            ]
            
            for field in content_fields:
                if field in item and item[field]:
                    content_parts.append(f"{field.title()}: {item[field]}")
            
            # Handle options/choices
            if 'options' in item and isinstance(item['options'], list):
                options_text = "Options: " + "; ".join([
                    f"{opt.get('id', i)}: {opt.get('text', opt)}" 
                    for i, opt in enumerate(item['options'])
                ])
                content_parts.append(options_text)
            
            return '\n'.join(content_parts)
        
        return str(item)
    
    def load_directory(
        self, 
        directory_path: str, 
        recursive: bool = True,
        default_metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """
        Load all supported documents from a directory
        
        Args:
            directory_path: Path to the directory
            recursive: Whether to search recursively
            default_metadata: Default metadata for all documents
            
        Returns:
            List of Document objects
        """
        directory_path = Path(directory_path)
        
        if not directory_path.exists() or not directory_path.is_dir():
            logger.error(f"Directory not found: {directory_path}")
            return []
        
        all_documents = []
        
        # Get all files with supported extensions
        pattern = "**/*" if recursive else "*"
        
        for ext in self.supported_extensions:
            files = list(directory_path.glob(f"{pattern}{ext}"))
            
            for file_path in files:
                # Create metadata for this file
                file_metadata = default_metadata.copy() if default_metadata else {}
                file_metadata.update({
                    'directory': str(directory_path),
                    'relative_path': str(file_path.relative_to(directory_path))
                })
                
                documents = self.load_document(file_path, file_metadata)
                all_documents.extend(documents)
        
        logger.info(f"✅ Loaded {len(all_documents)} total documents from {directory_path}")
        return all_documents
    
    def ingest_documents(self, documents: List[Document], subject: str = None) -> bool:
        """
        Ingest documents into the RAG system for a specific subject
        
        Args:
            documents: List of Document objects to ingest
            subject: Subject to ingest documents into (optional, uses default if None)
            
        Returns:
            Success status
        """
        if not documents:
            logger.warning("No documents provided for ingestion")
            return False
        
        try:
            success = rag_service.add_documents(documents, subject)
            if success:
                subject_info = f" for subject '{subject}'" if subject else ""
                logger.info(f"✅ Successfully ingested {len(documents)} documents{subject_info}")
            else:
                logger.error(f"❌ Failed to ingest documents for subject '{subject}'")
            
            return success
            
        except Exception as e:
            logger.error(f"❌ Error during document ingestion for subject '{subject}': {e}")
            return False
    
    def ingest_from_file(
        self,
        file_path: str,
        metadata: Optional[Dict[str, Any]] = None,
        subject: str = None
    ) -> bool:
        """
        Load and ingest a single file for a specific subject
        
        Args:
            file_path: Path to the file
            metadata: Additional metadata
            subject: Subject to ingest documents into (optional, uses default if None)
            
        Returns:
            Success status
        """
        documents = self.load_document(file_path, metadata)
        return self.ingest_documents(documents, subject)
    
    def ingest_from_directory(
        self,
        directory_path: str,
        recursive: bool = True,
        default_metadata: Optional[Dict[str, Any]] = None,
        subject: str = None
    ) -> bool:
        """
        Load and ingest all documents from a directory for a specific subject
        
        Args:
            directory_path: Path to the directory
            recursive: Whether to search recursively
            default_metadata: Default metadata for all documents
            subject: Subject to ingest documents into (optional, uses default if None)
            
        Returns:
            Success status
        """
        documents = self.load_directory(directory_path, recursive, default_metadata)
        return self.ingest_documents(documents, subject)
    
    def ingest_subject_from_metadata(self, documents: List[Document]) -> Dict[str, bool]:
        """
        Ingest documents into subject-specific collections based on their metadata
        
        Args:
            documents: List of Document objects with subject metadata
            
        Returns:
            Dictionary mapping subjects to success status
        """
        if not documents:
            logger.warning("No documents provided for subject-based ingestion")
            return {}
        
        # Group documents by subject
        subject_groups = {}
        for doc in documents:
            subject = doc.metadata.get('subject', 'general')
            if subject not in subject_groups:
                subject_groups[subject] = []
            subject_groups[subject].append(doc)
        
        # Ingest each subject group
        results = {}
        for subject, subject_docs in subject_groups.items():
            try:
                success = self.ingest_documents(subject_docs, subject)
                results[subject] = success
                logger.info(f"✅ Ingested {len(subject_docs)} documents for subject '{subject}': {'Success' if success else 'Failed'}")
            except Exception as e:
                logger.error(f"❌ Failed to ingest documents for subject '{subject}': {e}")
                results[subject] = False
        
        return results

# Global ingestion pipeline instance
ingestion_pipeline = DataIngestionPipeline()